from flask import Flask, request, send_file,make_response,jsonify
from docx import Document
from docx2pdf import convert
import sqlite3
import os
from flask_cors import CORS
import pythoncom

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes
# Define the upload folder
UPLOAD_FOLDER = 'D:\Satheesh\Projects\python'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/fill_and_generate_pdf', methods=['POST'])
def fill_and_generate_pdf():
    # Initialize the COM library
    pythoncom.CoInitialize()
    print("Entry@@@@")
    if "file" not in request.files:
        return make_response(jsonify({"message": "No file uploaded"}), 400)
    
    file = request.files["file"]

    # Check if the file is allowed
    if file.filename == "":
        return make_response(jsonify({"message": "No selected file"}), 400)
    
    # Save the file to the specified directory
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], file.filename))

    print("Entry@@@@",file.filename)
    # Get data from SQLite database
    conn = sqlite3.connect('uta_docufy.db')
    cursor = conn.cursor()


    # cursor.execute('''CREATE TABLE IF NOT EXISTS students (
    #                  id INTEGER PRIMARY KEY,
    #                  firstname TEXT NOT NULL,
    #                 lastname TEXT NOT NULL,
    #                 age INTEGER
    #              )''')
    
    #Define an SQL statement to insert data into a table
    # sql_insert = """INSERT INTO students (id, firstname, lastname,age) 
    #         VALUES (?, ?, ?, ?)"""
    #     # Define data to be inserted
    # student_data = (7, 'Liu','Yonghe', 35)
    #     # Execute the SQL statement
    # cursor.execute(sql_insert, student_data)
    
    # conn.commit()
    cursor.execute("SELECT * FROM students where emailid='lie.yonghe@uta.edu'")
    data = cursor.fetchall()

    conn.close()

    
    # Create Word document
    doc = Document(UPLOAD_FOLDER+"\\"+file.filename)

    # Replace keys in the Word document with SQLite3 query response
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '@surname@' in cell.text:  # Replace 'key' with your actual key
                    cell.text = cell.text.replace('@surname@', str(data[0][2]))
                if '@firstname@' in cell.text:  # Replace 'key' with your actual key
                    cell.text = cell.text.replace('@firstname@', str(data[0][1]))
                if '@emailld@' in cell.text:  # Replace 'key' with your actual key
                    cell.text = cell.text.replace('@emailld@', str(data[0][3]))    
    
    
    # Iterate over paragraphs
    for paragraph in doc.paragraphs:
        # Check if the placeholder is in the paragraph text
        if '@surname@' in paragraph.text:
            # Split the paragraph text based on the placeholder
            parts = paragraph.text.split('@surname@')
            # Replace only the placeholder with the desired value
            paragraph.text = str(data[0][2]).join(parts)
        if '@firstname@' in paragraph.text:
            # Split the paragraph text based on the placeholder
            parts = paragraph.text.split('@firstname@')
            # Replace only the placeholder with the desired value
            paragraph.text = str(data[0][1]).join(parts)
        if '@emailld@' in paragraph.text:
            # Split the paragraph text based on the placeholder
            parts = paragraph.text.split('@emailld@')
            # Replace only the placeholder with the desired value
            paragraph.text = str(data[0][3]).join(parts)    
        

    # Save filled document
    doc.save(UPLOAD_FOLDER+'\milled_document.docx')
    docx_path =UPLOAD_FOLDER+"\milled_document.docx"
    pdf_path=UPLOAD_FOLDER+"\generated_pdf.pdf"
    # Convert filled document to PDF
    # Example: Convert .docx to .pdf using a separate library like python-docx-pdf

    # Return the PDF file
    #return send_file('D:\Satheesh\Projects\python\generated_pdf.pdf', as_attachment=True)
    print("^^^^^^^^^^6")
    convert(docx_path, pdf_path)
    print("Hello")
    ##return send_file(directory=os.path.dirname(pdf_path), filename=os.path.basename(pdf_path))

    return jsonify({"message": "File uploaded successfully"})
    
    #return ""

# Define a route to handle the POST request
@app.route('/insert-student-data', methods=['POST'])
def post_data():
    # Access JSON data from the request
    json_data = request.get_json()

    # Process the JSON data
    # For example, if JSON data contains 'key1' and 'key2'
    firstname = json_data.get('firstname')
    lastname = json_data.get('lastname')
    emailid = json_data.get('emailid')
    age = json_data.get('age')
    print("#######",emailid)

    conn = sqlite3.connect('uta_docufy.db')
    cursor = conn.cursor()

    #cursor.execute('''Drop table students''')
    
    # cursor.execute('''CREATE TABLE IF NOT EXISTS students (
    #                  id INTEGER PRIMARY KEY,
    #                  firstname TEXT NOT NULL,
    #                 lastname TEXT NOT NULL,
    #                emailid TEXT NOT NULL,
    #                 age INTEGER
    #              )''')

    #Define an SQL statement to insert data into a table
    sql_insert = """INSERT INTO students (firstname, lastname, age, emailid) 
            VALUES (?, ?, ?, ?)"""
        # Define data to be inserted
    student_data = (firstname, lastname, age, emailid)
        # Execute the SQL statement
    cursor.execute(sql_insert, student_data)
    
    conn.commit()

    # Perform operations based on the received data

    # Return a response, usually in JSON format
    return jsonify({'message': 'Data received successfully', 'firstname': firstname, 'lastname': lastname})


if __name__ == '__main__':
    app.run(debug=True)
